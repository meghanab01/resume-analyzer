import gradio as gr
import google.generativeai as genai
import io
import re
import os
from PyPDF2 import PdfReader
import pdfplumber
from docx import Document
from dotenv import load_dotenv

load_dotenv()
import google.generativeai as genai

api_key = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=api_key)


# Google Model
model = genai.GenerativeModel('gemini-2.5-flash')

def extract_text_from_pdf(file_path):
    """Extracts text from PDF while preserving links and handling complex layouts."""
    text = ""
    linkedin_links = []

    try:
        with open(file_path, 'rb') as f:
            pdf_bytes = f.read()

        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"

                    # Extract LinkedIn links from text
                    found_links = re.findall(r"https?://(www\.)?linkedin\.com/in/[A-Za-z0-9\-_/]+", page_text)
                    linkedin_links.extend(found_links)

                    # Check hyperlink metadata (fallback for pdfplumber)
                    if page.hyperlinks:
                        for link in page.hyperlinks:
                            uri = link.get("uri", "")
                            if "linkedin.com/in/" in uri:
                                linkedin_links.append(uri)

        except Exception as e:
            print(f"pdfplumber failed: {e}, falling back to PyPDF2")
            try:
                reader = PdfReader(io.BytesIO(pdf_bytes))
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text += page_text
                    found_links = re.findall(r"https?://(www\.)?linkedin\.com/in/[A-Za-z0-9\-_/]+", page_text)
                    linkedin_links.extend(found_links)
            except Exception as e:
                return f"PDF extraction failed: {str(e)}", []
    except Exception as e:
        return f"File error: {str(e)}", []
    
    return text.strip(), list(set(linkedin_links)) # Return text and unique links

def extract_text_from_docx(file_path):
    """Extracts text from DOCX files."""
    text = ""
    linkedin_links = []
    try:
        document = Document(file_path)
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
            found_links = re.findall(r"https?://(www\.)?linkedin\.com/in/[A-Za-z0-9\-_/]+", paragraph.text)
            linkedin_links.extend(found_links)
        
        # Extracting info from tables in the files
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text or ""
                    text += cell_text + "\n"
                    found_links = re.findall(r"https?://(www\.)?linkedin\.com/in/[A-Za-z0-9\-_/]+", cell_text)
                    linkedin_links.extend(found_links)

    except Exception as e:
        return f"DOCX extraction failed: {str(e)}", []
    
    return text.strip(), list(set(linkedin_links)) # Return text and unique links

def extract_text_from_file(file):
    """Determines file type and extracts text accordingly."""
    if file is None:
        return "Error: No file uploaded.", []

    file_path = file.name
    file_extension = os.path.splitext(file_path)[1].lower()

    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension == ".docx":
        return extract_text_from_docx(file_path)
    else:
        return "Error: Unsupported file type. Please upload a PDF or DOCX.", []

def clean_text(text):
    """Clean excessive whitespace while keeping links."""
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def analyze_resume(file, job_description):
    """Analyzes resume against job description using Gemini."""
    resume_text, linkedin_links = extract_text_from_file(file)

    if resume_text.startswith("Error:"):
        return f"⚠️ {resume_text}"
    
    clean_resume = clean_text(resume_text)
   
    if linkedin_links and "📎 LinkedIn Profiles:" not in clean_resume:
        clean_resume += "\n\n📎 LinkedIn Profiles:\n" + "\n".join(f"- {link}" for link in linkedin_links)

    prompt = f"""You're an AI-powered resume evaluator.

**Resume Content (includes LinkedIn if found):**
{clean_resume[:15000]}

**Job Description:**
{job_description[:5000]}

Provide the following structured analysis:

## 🔍 Match Score (0-100)
- Score and reasoning

## ✅ Top 3 Strengths
- Clear bullet points

## ❌ Gaps / Missing Skills
- Compared to job description

## 💡 Suggestions for Improvement
- Short and actionable

Make sure to preserve and highlight any LinkedIn profiles included in the resume.
"""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"⚠️ Analysis failed: {str(e)}"

# Gradio UI
with gr.Blocks(theme=gr.themes.Soft()) as demo:
    gr.Markdown("# 🧑‍💻 Universal Resume Analyzer")
    
    with gr.Row():
        with gr.Column():
            gr.Markdown("### Upload Resume")
            file_input = gr.File(
                label="Upload Resume (PDF or DOCX, supports LinkedIn detection)",
                file_types=[".pdf", ".docx"],
                type="filepath"
            )
            jd_input = gr.Textbox(
                label="Paste Job Description",
                lines=7,
                # Placeholder text
                placeholder="Describe the role, responsibilities, and required skills for the job. For example: 'This role requires strong leadership, project management, and communication skills. Experience in marketing and client relations is a plus.'"
            )
            btn = gr.Button("Analyze", variant="primary")
            
        with gr.Column():
            gr.Markdown("### Analysis Results")
            output = gr.Markdown(label="Feedback", latex_delimiters=[])
    
    btn.click(
        fn=analyze_resume,
        inputs=[file_input, jd_input],
        outputs=output
    )

if __name__ == "__main__":
    demo.launch(share=True)

